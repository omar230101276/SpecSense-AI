import docx
import io
import numpy as np
import cv2
from PIL import Image

def process_docx(docx_path, ocr_engine, detail=1):
    """
    Extract text and images from a DOCX file.
    
    :param docx_path: Path to the .docx file.
    :param ocr_engine: Instance of OCREngine to process embedded images.
    :param detail: Detail level for OCR results.
    :return: List of results in EasyOCR format: [([[x,y]..], text, prob), ...]
    """
    results = []
    
    try:
        doc = docx.Document(docx_path)
        
        # 1. Extract Text from Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                # Create a "fake" OCR result for valid text
                # Bbox is dummy [[0,0], [0,0], [0,0], [0,0]]
                # Confidence is 1.0 because it's digital text
                results.append(([[0,0], [1,0], [1,1], [0,1]], para.text.strip(), 1.0))
                
        # 2. Extract Text from Tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                         results.append(([[0,0], [1,0], [1,1], [0,1]], cell.text.strip(), 1.0))

        # 3. Extract Images (Advanced)
        # Iterate through relationships to find image parts
        # This covers images embedded in the document
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                try:
                    image_data = rel.target_part.blob
                    # Convert bytes to numpy array
                    # using PIL first to handle formats safely
                    pil_img = Image.open(io.BytesIO(image_data))
                    pil_img = pil_img.convert('RGB')
                    img_array = np.array(pil_img)
                    
                    print(f"Found embedded image of size {img_array.shape}, running OCR...")
                    ocr_results = ocr_engine.read_image_from_array(img_array, detail=detail)
                    results.extend(ocr_results)
                except Exception as img_e:
                    print(f"Failed to process an embedded image: {img_e}")

    except Exception as e:
        print(f"Error processing DOCX: {e}")
        return []

    # Filter results based on detail level
    if detail == 0:
        # Return only the text string
        return [r[1] for r in results]
    else:
        return results
